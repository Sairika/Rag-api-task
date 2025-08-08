import os
import fitz  # PyMuPDF
from docx import Document
import pandas as pd
import sqlite3
from PIL import Image
import pytesseract
from typing import List, Dict, Any
import uuid
from pathlib import Path
import re

class DocumentProcessor:
    def __init__(self, chunk_size: int = 800, chunk_overlap: int = 100):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        print(f"📝 Document processor initialized (chunk_size: {chunk_size}, overlap: {chunk_overlap})")
    
    def process_document(self, file_path: str, filename: str) -> List[Dict[str, Any]]:
        """Process a document and return text chunks with metadata"""
        
        file_ext = Path(file_path).suffix.lower()
        print(f"🔄 Processing {filename} ({file_ext})...")
        
        try:
            if file_ext == '.pdf':
                chunks = self._process_pdf(file_path, filename)
            elif file_ext == '.docx':
                chunks = self._process_docx(file_path, filename)
            elif file_ext == '.txt':
                chunks = self._process_txt(file_path, filename)
            elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
                chunks = self._process_image(file_path, filename)
            elif file_ext == '.csv':
                chunks = self._process_csv(file_path, filename)
            elif file_ext == '.db':
                chunks = self._process_sqlite(file_path, filename)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            print(f"✅ Processed {filename}: {len(chunks)} chunks created")
            return chunks
                
        except Exception as e:
            print(f"❌ Error processing {filename}: {e}")
            return []
    
    def _process_pdf(self, file_path: str, filename: str) -> List[Dict[str, Any]]:
        """Process PDF file with better text extraction"""
        chunks = []
        
        try:
            doc = fitz.open(file_path)
            print(f"📄 PDF has {len(doc)} pages")
            
            for page_num, page in enumerate(doc, 1):
                print(f"🔄 Processing page {page_num}/{len(doc)}")
                
                # Extract text
                text = page.get_text()
                
                # If text is minimal, try OCR on the page
                if len(text.strip()) < 50:
                    try:
                        # Convert page to image and OCR
                        pix = page.get_pixmap()
                        img_data = pix.tobytes("png")
                        
                        # Use PIL and OCR
                        from io import BytesIO
                        img = Image.open(BytesIO(img_data))
                        ocr_text = self.extract_text_from_image(img)
                        
                        if len(ocr_text.strip()) > len(text.strip()):
                            text = ocr_text
                            print(f"📷 Used OCR for page {page_num} (better results)")
                    
                    except Exception as e:
                        print(f"⚠️ OCR failed for page {page_num}: {e}")
                
                if text.strip():
                    # Clean up the text
                    text = self._clean_text(text)
                    
                    page_chunks = self._create_chunks(text, {
                        'filename': filename,
                        'page': page_num,
                        'file_type': 'pdf'
                    })
                    chunks.extend(page_chunks)
            
            doc.close()
            
        except Exception as e:
            print(f"❌ Error processing PDF {filename}: {e}")
        
        return chunks
    
    def _process_docx(self, file_path: str, filename: str) -> List[Dict[str, Any]]:
        """Process Word document with better structure handling"""
        chunks = []
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs with structure preservation
            text_content = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    # Preserve heading structure
                    if para.style.name.startswith('Heading'):
                        text_content.append(f"\n### {para.text.strip()}\n")
                    else:
                        text_content.append(para.text.strip())
            
            # Extract text from tables with better formatting
            for table in doc.tables:
                table_text = ["[TABLE START]"]
                
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_cells.append(cell_text)
                    
                    if row_cells:
                        table_text.append(" | ".join(row_cells))
                
                table_text.append("[TABLE END]")
                text_content.extend(table_text)
            
            # Join all content
            full_text = "\n".join(text_content)
            
            if full_text.strip():
                # Clean and process
                full_text = self._clean_text(full_text)
                chunks = self._create_chunks(full_text, {
                    'filename': filename,
                    'page': 1,
                    'file_type': 'docx'
                })
        
        except Exception as e:
            print(f"❌ Error processing DOCX {filename}: {e}")
        
        return chunks
    
    def _process_txt(self, file_path: str, filename: str) -> List[Dict[str, Any]]:
        """Process text file with encoding detection"""
        chunks = []
        
        # Try multiple encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1', 'utf-16']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                
                if text.strip():
                    print(f"✅ Successfully read with {encoding} encoding")
                    text = self._clean_text(text)
                    chunks = self._create_chunks(text, {
                        'filename': filename,
                        'page': 1,
                        'file_type': 'txt'
                    })
                break
                
            except UnicodeDecodeError:
                continue
            except Exception as e:
                print(f"❌ Error with encoding {encoding}: {e}")
                continue
        
        if not chunks:
            print(f"⚠️ Could not read {filename} with any encoding")
        
        return chunks
    
    def _process_image(self, file_path: str, filename: str) -> List[Dict[str, Any]]:
        """Process image file using OCR with preprocessing"""
        chunks = []
        
        try:
            print(f"🖼️ Processing image with OCR...")
            image = Image.open(file_path)
            
            # Preprocess image for better OCR
            image = self._preprocess_image_for_ocr(image)
            
            text = self.extract_text_from_image(image)
            
            if text.strip():
                text = self._clean_text(text)
                chunks = self._create_chunks(text, {
                    'filename': filename,
                    'page': 1,
                    'file_type': 'image_ocr'
                })
                print(f"📝 Extracted {len(text)} characters from image")
            else:
                print("⚠️ No text found in image")
        
        except Exception as e:
            print(f"❌ Error processing image {filename}: {e}")
        
        return chunks
    
    def _process_csv(self, file_path: str, filename: str) -> List[Dict[str, Any]]:
        """Process CSV file with comprehensive analysis"""
        chunks = []
        
        try:
            # Try different separators and encodings
            separators = [',', ';', '\t']
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            df = None
            for sep in separators:
                for enc in encodings:
                    try:
                        df = pd.read_csv(file_path, sep=sep, encoding=enc)
                        if len(df.columns) > 1:  # Successful parsing likely has multiple columns
                            print(f"✅ CSV parsed with separator '{sep}' and encoding '{enc}'")
                            break
                    except:
                        continue
                if df is not None and len(df.columns) > 1:
                    break
            
            if df is None or df.empty:
                print(f"⚠️ Could not parse CSV file {filename}")
                return []
            
            # Convert DataFrame to comprehensive text representation
            text_parts = []
            
            # Add header information
            headers = list(df.columns)
            text_parts.append(f"CSV FILE: {filename}")
            text_parts.append(f"COLUMNS ({len(headers)}): {' | '.join(headers)}")
            text_parts.append(f"TOTAL ROWS: {len(df)}")
            
            # Add data types
            text_parts.append("\nCOLUMN TYPES:")
            for col in df.columns:
                dtype = str(df[col].dtype)
                text_parts.append(f"- {col}: {dtype}")
            
            # Add sample data (first 10 rows)
            text_parts.append(f"\nSAMPLE DATA (first {min(10, len(df))} rows):")
            for idx, row in df.head(10).iterrows():
                row_text = " | ".join([f"{col}: {str(val)}" for col, val in zip(df.columns, row.values)])
                text_parts.append(f"Row {idx + 1}: {row_text}")
            
            # Add statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
            if numeric_cols:
                text_parts.append("\nNUMERIC STATISTICS:")
                for col in numeric_cols:
                    try:
                        stats = df[col].describe()
                        text_parts.append(f"- {col}: mean={stats['mean']:.2f}, std={stats['std']:.2f}, min={stats['min']:.2f}, max={stats['max']:.2f}")
                    except:
                        continue
            
            # Add unique values for categorical columns (if reasonable number)
            categorical_cols = df.select_dtypes(include=['object']).columns.tolist()
            if categorical_cols:
                text_parts.append("\nCATEGORICAL DATA:")
                for col in categorical_cols[:5]:  # Limit to first 5 categorical columns
                    try:
                        unique_vals = df[col].unique()
                        if len(unique_vals) <= 20:  # Only if manageable number
                            text_parts.append(f"- {col} unique values: {', '.join(map(str, unique_vals))}")
                        else:
                            text_parts.append(f"- {col}: {len(unique_vals)} unique values")
                    except:
                        continue
            
            full_text = "\n".join(text_parts)
            
            if full_text.strip():
                chunks = self._create_chunks(full_text, {
                    'filename': filename,
                    'page': 1,
                    'file_type': 'csv'
                })
        
        except Exception as e:
            print(f"❌ Error processing CSV {filename}: {e}")
        
        return chunks
    
    def _process_sqlite(self, file_path: str, filename: str) -> List[Dict[str, Any]]:
        """Process SQLite database file with comprehensive analysis"""
        chunks = []
        
        try:
            conn = sqlite3.connect(file_path)
            cursor = conn.cursor()
            
            # Get table names
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
            tables = cursor.fetchall()
            
            if not tables:
                print(f"⚠️ No tables found in database {filename}")
                return []
            
            text_parts = [f"SQLite DATABASE: {filename}"]
            text_parts.append(f"CONTAINS {len(tables)} TABLES: {[t[0] for t in tables]}")
            
            for table in tables:
                table_name = table[0]
                
                try:
                    # Get table schema
                    cursor.execute(f"PRAGMA table_info({table_name})")
                    columns = cursor.fetchall()
                    column_info = [(col[1], col[2]) for col in columns]  # name, type
                    
                    text_parts.append(f"\nTABLE: {table_name}")
                    text_parts.append(f"COLUMNS: {', '.join([f'{name}({dtype})' for name, dtype in column_info])}")
                    
                    # Get row count
                    cursor.execute(f"SELECT COUNT(*) FROM {table_name}")
                    row_count = cursor.fetchone()[0]
                    text_parts.append(f"ROWS: {row_count}")
                    
                    # Get sample data
                    if row_count > 0:
                        cursor.execute(f"SELECT * FROM {table_name} LIMIT 5")
                        sample_rows = cursor.fetchall()
                        
                        text_parts.append("SAMPLE DATA:")
                        column_names = [col[0] for col in column_info]
                        
                        for i, row in enumerate(sample_rows):
                            row_dict = dict(zip(column_names, row))
                            row_text = " | ".join([f"{k}: {v}" for k, v in row_dict.items()])
                            text_parts.append(f"  Row {i+1}: {row_text}")
                
                except Exception as e:
                    text_parts.append(f"  Error reading table {table_name}: {e}")
            
            conn.close()
            
            full_text = "\n".join(text_parts)
            
            if full_text.strip():
                chunks = self._create_chunks(full_text, {
                    'filename': filename,
                    'page': 1,
                    'file_type': 'sqlite'
                })
        
        except Exception as e:
            print(f"❌ Error processing SQLite {filename}: {e}")
        
        return chunks
    
    def extract_text_from_image(self, image: Image.Image) -> str:
        """Extract text from image using OCR with better preprocessing"""
        try:
            # Use pytesseract to extract text with better configuration
            custom_config = r'--oem 3 --psm 6 -c tessedit_char_whitelist=ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789.,!?;:()\'"- '
            
            text = pytesseract.image_to_string(
                image, 
                lang='eng',
                config=custom_config
            )
            
            return self._clean_text(text)
            
        except Exception as e:
            print(f"❌ OCR error: {e}")
            # Fallback to basic OCR
            try:
                text = pytesseract.image_to_string(image, lang='eng')
                return self._clean_text(text)
            except:
                return ""
    
    def _preprocess_image_for_ocr(self, image: Image.Image) -> Image.Image:
        """Preprocess image to improve OCR accuracy"""
        try:
            # Convert to RGB if needed
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Convert to grayscale for better OCR
            image = image.convert('L')
            
            # Resize if too small (OCR works better on larger images)
            width, height = image.size
            if width < 300 or height < 300:
                scale_factor = max(300/width, 300/height)
                new_size = (int(width * scale_factor), int(height * scale_factor))
                image = image.resize(new_size, Image.LANCZOS)
            
            return image
            
        except Exception as e:
            print(f"⚠️ Image preprocessing failed: {e}")
            return image
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)  # Multiple blank lines to double
        text = re.sub(r'[ \t]+', ' ', text)      # Multiple spaces/tabs to single space
        
        # Remove weird characters but keep useful punctuation
        text = re.sub(r'[^\w\s.,!?;:()\'"@#$%&*+=\[\]{}<>/\\|-]', '', text)
        
        # Fix common OCR mistakes
        text = text.replace('|', 'I')  # Common OCR error
        text = text.replace('0', 'O')  # When appropriate (context-dependent)
        
        return text.strip()
    
    def _create_chunks(self, text: str, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Split text into chunks with improved sentence-boundary awareness"""
        chunks = []
        text = text.strip()
        
        if not text:
            return chunks
        
        # First, split into sentences for better boundary detection
        sentences = self._split_into_sentences(text)
        
        if not sentences:
            return chunks
        
        current_chunk = ""
        chunk_index = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            # Calculate if adding this sentence would exceed chunk size
            potential_chunk = current_chunk + (" " if current_chunk else "") + sentence
            
            if len(potential_chunk) > self.chunk_size and current_chunk:
                # Save current chunk
                if current_chunk.strip():
                    chunk_metadata = metadata.copy()
                    chunk_metadata['chunk_index'] = chunk_index
                    chunk_metadata['chunk_id'] = f"{metadata['filename']}_{chunk_index}"
                    
                    chunks.append({
                        'content': current_chunk.strip(),
                        'metadata': chunk_metadata,
                        'chunk_id': chunk_metadata['chunk_id']
                    })
                    
                    chunk_index += 1
                
                # Start new chunk with overlap
                if self.chunk_overlap > 0 and current_chunk:
                    # Keep last few sentences for context
                    overlap_sentences = self._get_overlap_sentences(current_chunk, self.chunk_overlap)
                    current_chunk = overlap_sentences + " " + sentence
                else:
                    current_chunk = sentence
            else:
                current_chunk = potential_chunk
        
        # Add the final chunk
        if current_chunk.strip():
            chunk_metadata = metadata.copy()
            chunk_metadata['chunk_index'] = chunk_index
            chunk_metadata['chunk_id'] = f"{metadata['filename']}_{chunk_index}"
            
            chunks.append({
                'content': current_chunk.strip(),
                'metadata': chunk_metadata,
                'chunk_id': chunk_metadata['chunk_id']
            })
        
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences using multiple delimiters"""
        
        # Split on sentence endings, but be careful with abbreviations
        sentence_endings = r'(?<![A-Z][a-z]\.)\s*[.!?]+\s+'
        sentences = re.split(sentence_endings, text)
        
        # Clean up sentences
        cleaned_sentences = []
        for sentence in sentences:
            sentence = sentence.strip()
            if sentence and len(sentence) > 10:  # Ignore very short fragments
                cleaned_sentences.append(sentence)
        
        return cleaned_sentences
    
    def _get_overlap_sentences(self, text: str, max_overlap_chars: int) -> str:
        """Get the last few sentences that fit within overlap limit"""
        
        sentences = self._split_into_sentences(text)
        if not sentences:
            return ""
        
        # Start from the end and build overlap
        overlap_text = ""
        for sentence in reversed(sentences):
            potential_overlap = sentence + " " + overlap_text if overlap_text else sentence
            if len(potential_overlap) <= max_overlap_chars:
                overlap_text = potential_overlap
            else:
                break
        
        return overlap_text.strip()